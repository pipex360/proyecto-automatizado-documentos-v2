from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import shutil
import os
import re

class ProcesadorPlantillas:
    def __init__(self, ruta_plantilla):
        self.ruta_plantilla = ruta_plantilla
        self.doc = None
    
    def reemplazar_en_parrafo(self, parrafo, datos):
        """Reemplaza marcadores en un párrafo preservando formato"""
        texto_completo = parrafo.text
        
        if '{{' not in texto_completo or '}}' not in texto_completo:
            return
        
        # Hacer reemplazos
        texto_nuevo = texto_completo
        for key, value in datos.items():
            marcador = f"{{{{{key}}}}}"
            texto_nuevo = texto_nuevo.replace(marcador, str(value))
        
        if texto_nuevo != texto_completo:
            # Guardar formato del párrafo
            estilo_parrafo = parrafo.style
            alineacion = parrafo.alignment
            
            # Guardar formato del primer run si existe
            formato_run = None
            if parrafo.runs:
                primer_run = parrafo.runs[0]
                formato_run = {
                    'bold': primer_run.bold,
                    'italic': primer_run.italic,
                    'underline': primer_run.underline,
                    'font_name': primer_run.font.name,
                    'font_size': primer_run.font.size
                }
            
            # Limpiar runs
            for run in parrafo.runs:
                run.text = ''
            
            # Agregar texto nuevo
            if parrafo.runs:
                nuevo_run = parrafo.runs[0]
            else:
                nuevo_run = parrafo.add_run()
            
            nuevo_run.text = texto_nuevo
            
            # Restaurar formato del run
            if formato_run:
                if formato_run['bold'] is not None:
                    nuevo_run.bold = formato_run['bold']
                if formato_run['italic'] is not None:
                    nuevo_run.italic = formato_run['italic']
                if formato_run['underline'] is not None:
                    nuevo_run.underline = formato_run['underline']
                if formato_run['font_name']:
                    nuevo_run.font.name = formato_run['font_name']
                if formato_run['font_size']:
                    nuevo_run.font.size = formato_run['font_size']
            
            # Restaurar formato del párrafo
            parrafo.style = estilo_parrafo
            if alineacion:
                parrafo.alignment = alineacion
            else:
                # Si no tiene alineación definida, usar justificado
                parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def procesar(self, datos, ruta_salida):
        try:
            self.doc = Document(self.ruta_plantilla)
            
            # Reemplazar en párrafos
            for para in self.doc.paragraphs:
                self.reemplazar_en_parrafo(para, datos)
            
            # Reemplazar en tablas
            for tabla in self.doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            self.reemplazar_en_parrafo(para, datos)
            
            # Reemplazar en headers y footers
            for section in self.doc.sections:
                for para in section.header.paragraphs:
                    self.reemplazar_en_parrafo(para, datos)
                for para in section.footer.paragraphs:
                    self.reemplazar_en_parrafo(para, datos)
                
                for tabla in section.header.tables:
                    for fila in tabla.rows:
                        for celda in fila.cells:
                            for para in celda.paragraphs:
                                self.reemplazar_en_parrafo(para, datos)
                
                for tabla in section.footer.tables:
                    for fila in tabla.rows:
                        for celda in fila.cells:
                            for para in celda.paragraphs:
                                self.reemplazar_en_parrafo(para, datos)
            
            # Guardar
            self.doc.save(ruta_salida)
            
            # Segundo paso: XML
            temp_salida = ruta_salida + ".temp"
            if self.reemplazar_en_xml(ruta_salida, datos, temp_salida):
                shutil.move(temp_salida, ruta_salida)
            
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False
    
    def reemplazar_en_xml(self, ruta_docx, datos, ruta_salida):
        try:
            temp_dir = 'temp_xml'
            os.makedirs(temp_dir, exist_ok=True)
            
            with zipfile.ZipFile(ruta_docx, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            for carpeta, subcarpetas, archivos in os.walk(temp_dir):
                for archivo in archivos:
                    if archivo.endswith('.xml'):
                        ruta = os.path.join(carpeta, archivo)
                        with open(ruta, 'r', encoding='utf-8') as f:
                            contenido = f.read()
                        
                        contenido_nuevo = contenido
                        for key, value in datos.items():
                            marcador = f"{{{{{key}}}}}"
                            contenido_nuevo = contenido_nuevo.replace(marcador, str(value))
                        
                        if contenido_nuevo != contenido:
                            with open(ruta, 'w', encoding='utf-8') as f:
                                f.write(contenido_nuevo)
            
            with zipfile.ZipFile(ruta_salida, 'w', zipfile.ZIP_DEFLATED) as docx:
                for carpeta, subcarpetas, archivos in os.walk(temp_dir):
                    for archivo in archivos:
                        ruta_completa = os.path.join(carpeta, archivo)
                        ruta_en_zip = os.path.relpath(ruta_completa, temp_dir)
                        docx.write(ruta_completa, ruta_en_zip)
            
            shutil.rmtree(temp_dir)
            return True
        except Exception as e:
            print(f"Error XML: {e}")
            return False
